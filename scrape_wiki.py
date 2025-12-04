from playwright.sync_api import sync_playwright
from docx import Document
from urllib.parse import urljoin, urlparse, urlunparse
import json
import os

# ----------------------------
# Configuration
# ----------------------------
START_URL = "https://sites.google.com/view/kopievonfitreisengroupwiki/fit-group-wiki"
AUTH_FILE = "auth.json"

OUTPUT_DIR = "C:/Users/saq/Downloads/wiki_export"
os.makedirs(OUTPUT_DIR, exist_ok=True)

OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "wiki_export.docx")
STATE_FILE = os.path.join(OUTPUT_DIR, "crawler_state.json")  # save visited + queue
SAVE_INTERVAL = 5  # save every 5 pages

# ----------------------------
# Utility functions
# ----------------------------
def normalize_url(url):
    """Remove query params and fragments for duplicate detection."""
    parsed = urlparse(url)
    normalized = parsed._replace(query="", fragment="")
    return urlunparse(normalized)

def save_state(doc, visited, queue):
    """Save Word document + crawler state to disk."""
    doc.save(OUTPUT_DOCX)
    state = {
        "visited": list(visited),
        "queue": list(queue)
    }
    with open(STATE_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2)
    print(f"💾 Progress saved ({len(visited)} pages, {len(queue)} in queue)")

def load_state():
    """Load crawler state + Word document."""
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, "r", encoding="utf-8") as f:
            state = json.load(f)
        visited = set(state.get("visited", []))
        queue = state.get("queue", [])
        print(f"🔄 Loaded state: {len(visited)} visited, {len(queue)} in queue")
    else:
        visited = set()
        queue = [START_URL]
        print("📝 Starting fresh crawl")
    if os.path.exists(OUTPUT_DOCX):
        doc = Document(OUTPUT_DOCX)
        print("📄 Existing document loaded")
    else:
        doc = Document()
        doc.add_heading("Company Wiki Export", level=0)
        print("📝 New document created")
    return doc, visited, queue

# ----------------------------
# Scrape a single page
# ----------------------------
def scrape_page(page, url):
    print(f"➡ Visiting: {url}")
    try:
        page.goto(url, timeout=60000)
        page.wait_for_load_state("domcontentloaded")
    except:
        print("⚠ Timeout loading page, skipping")
        return []

    title = page.title() or "Untitled Page"
    text = page.text_content("body") or ""

    # Save page content
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Source: {url}\n")
    doc.add_paragraph(text)

    # Extract internal wiki links
    new_links = []
    anchors = page.locator("a").all()
    for a in anchors:
        try:
            href = a.get_attribute("href")
            if not href:
                continue
            full_url = urljoin(url, href)
            # Only follow Google Sites / your wiki
            if "sites.google.com/view/kopievonfitreisengroupwiki" in full_url:
                norm_url = normalize_url(full_url)
                if norm_url not in visited and norm_url not in queue:
                    new_links.append(norm_url)
        except:
            pass
    return new_links

# ----------------------------
# Main crawler
# ----------------------------
if __name__ == "__main__":
    doc, visited, queue = load_state()
    pages_crawled = 0

    with sync_playwright() as p:
        browser = p.chromium.launch(
            channel="msedge",
            headless=True,
            args=["--disable-blink-features=AutomationControlled"]
        )
        context = browser.new_context(storage_state=AUTH_FILE)
        page = context.new_page()

        while queue:
            url = queue.pop(0)
            norm_url = normalize_url(url)
            if norm_url in visited:
                continue

            visited.add(norm_url)
            new_links = scrape_page(page, url)
            queue.extend(new_links)
            pages_crawled += 1

            if pages_crawled % SAVE_INTERVAL == 0:
                save_state(doc, visited, queue)

        browser.close()

    # Final save
    save_state(doc, visited, queue)
    print("\n🎉 Wiki export completed!")
    print(f"📁 Final file saved at: {OUTPUT_DOCX}")

import re
from docx import Document

# Input and output files
INPUT_FILE = r"C:\Users\saq\Downloads\wiki_export\wiki_export.docx"
OUTPUT_FILE = r"C:\Users\saq\Downloads\wiki_export\extracted_links.docx"

# Regex to match only the real Google Sites links
OLD_DOMAIN_PATTERN = r"https://sites\.google\.com/view/kopievonfitreisengroupwiki/[\w\-/]+"

def extract_links_with_titles(file_path):
    doc = Document(file_path)
    links = []
    previous_para = ""
    
    for para in doc.paragraphs:
        matches = re.findall(OLD_DOMAIN_PATTERN, para.text)
        for link in matches:
            # Use the previous paragraph as the title
            title = previous_para.strip()
            # Replace old domain with new domain
            new_link = link.replace(
                "https://sites.google.com/view/kopievonfitreisengroupwiki",
                "https://sites.google.com/site/fitgroupwiki2"
            )
            links.append((title, new_link))
        previous_para = para.text  # store current paragraph for next iteration
    return links

# Create new docx for results
new_doc = Document()
extracted_links = extract_links_with_titles(INPUT_FILE)

for title, link in extracted_links:
    new_doc.add_paragraph(f"Title: {title}")
    new_doc.add_paragraph(f"Source: {link}")
    new_doc.add_paragraph("")  # spacing

# Save document
new_doc.save(OUTPUT_FILE)
print(f"Extraction complete! Saved to: {OUTPUT_FILE}")

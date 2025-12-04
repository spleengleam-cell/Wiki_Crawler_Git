import re
from docx import Document
from docx.document import Document as _Document

NOISE_KEYWORDS = [
    "oauth", "osid", "pageid", "palette_colors", "passwdsession",
    "platform_variant", "pstmsg", "pav"
]

GARBAGE_PATTERNS = [
    r"u00[0-9A-Fa-f]{2}",
    r"\\u[0-9A-Fa-f]{4}",
    r"[A-Za-z]\s*=\s*[A-Za-z0-9]+",
    r"P=function\(.*?\}",
    r"[{}()\\/\[\];]{2,}",
    r"[A-Za-z0-9_$]{15,}",
]

def clean_text(text):
    for k in NOISE_KEYWORDS:
        text = re.sub(rf"\b{k}\b", "", text, flags=re.IGNORECASE)
    for pat in GARBAGE_PATTERNS:
        text = re.sub(pat, " ", text)
    return re.sub(r"\s+", " ", text).strip()


doc = Document(r"C:\Users\saq\Downloads\wiki_export\wiki_export.docx")
new_doc = Document()

for para in doc.paragraphs:
    cleaned = clean_text(para.text)
    if cleaned:
        new_doc.add_paragraph(cleaned)

new_doc.save(r"C:\Users\saq\Downloads\wiki_export\cleaned_wiki.docx")
print("Saved → cleaned_output.docx")
